import os
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def add_ismac_standard_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    
    p_pr = footer_p._element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    top_bdr = OxmlElement('w:top')
    top_bdr.set(qn('w:val'), 'single')
    top_bdr.set(qn('w:sz'), '6')
    top_bdr.set(qn('w:space'), '4')
    top_bdr.set(qn('w:color'), '000000')
    p_bdr.append(top_bdr)
    p_pr.append(p_bdr)
    
    table = footer.add_table(rows=1, cols=3, width=Inches(7.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c0 = table.cell(0, 0)
    c0.width = Inches(2.2)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run("Tél. : 212 (0)5 37 27 17 00\nFax : 212 (0)5 37 77 38 65")
    r0.font.size = Pt(8.5)
    r0.font.name = 'Arial'
    
    c1 = table.cell(0, 1)
    c1.width = Inches(3.2)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("ISMAC - Av. Allal El Fassi, Cité Al Irfane B.P 6598 - Souissi, Rabat/Maroc\nشارع علال الفاسي - مدينة العرفان السويسي - الرباط - المغرب - ص.ب 6598")
    r1.font.size = Pt(7.5)
    r1.font.name = 'Arial'
    
    c2 = table.cell(0, 2)
    c2.width = Inches(1.8)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("E-mail : info@ismac.ac.ma\nwww.ismac.ac.ma")
    r2.font.size = Pt(8.5)
    r2.font.name = 'Arial'

# 1. BON DE DECHARGE FOURNITURE INFORMATIQUE
def generate_decharge_informatique(beneficiary, date_str, items):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_t = doc.add_paragraph()
    r_t = p_t.add_run("BON DE DÉCHARGE DE FOURNITURE INFORMATIQUE")
    r_t.bold = True
    r_t.font.size = Pt(14)
    r_t.font.name = 'Arial'
    p_t.paragraph_format.space_after = Pt(24)
    
    p_rec = doc.add_paragraph()
    r_l1 = p_rec.add_run("RÉCUPÉRÉ PAR  : ")
    r_l1.bold = True
    r_l1.font.size = Pt(11)
    r_v1 = p_rec.add_run(f"{beneficiary.upper()}\t\t")
    r_v1.font.size = Pt(11)
    p_rec.paragraph_format.space_after = Pt(4)
    
    p_date = doc.add_paragraph()
    r_l2 = p_date.add_run("DATE DE SORTIE : ")
    r_l2.bold = True
    r_l2.font.size = Pt(11)
    r_v2 = p_date.add_run(date_str)
    r_v2.font.size = Pt(11)
    p_date.paragraph_format.space_after = Pt(20)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_w = [Inches(4.2), Inches(1.1), Inches(1.7)]
    headers = ["Matériel", "Quantité", "Observations"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = col_w[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].bold = True
        set_cell_background(hdr_cells[i], "EFEFEF")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
    for it in items:
        row = table.add_row().cells
        row[0].text = it.get('name', '')
        row[0].width = col_w[0]
        set_cell_margins(row[0], top=100, bottom=100, left=150, right=150)
        
        row[1].text = str(it.get('quantity', 1))
        row[1].width = col_w[1]
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row[1], top=100, bottom=100, left=150, right=150)
        
        row[2].text = it.get('observations') or "—-----------------"
        row[2].width = col_w[2]
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row[2], top=100, bottom=100, left=150, right=150)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(45)
    
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sig = p_sig.add_run("La personne concernée")
    r_sig.bold = True
    r_sig.font.size = Pt(11)
    
    add_ismac_standard_footer(doc)
    return doc

# 2. BON DE DECHARGE MATERIEL AUDIOVISUEL
def generate_decharge_audiovisuel(beneficiary, date_str, items):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_t = doc.add_paragraph()
    r_t = p_t.add_run("BON DE DÉCHARGE DE MATÉRIEL AUDIOVISUEL")
    r_t.bold = True
    r_t.font.size = Pt(14)
    r_t.font.name = 'Arial'
    p_t.paragraph_format.space_after = Pt(24)
    
    p_rec = doc.add_paragraph()
    r_l1 = p_rec.add_run("RÉCUPÉRÉ PAR  : ")
    r_l1.bold = True
    r_l1.font.size = Pt(11)
    r_v1 = p_rec.add_run(beneficiary.upper())
    r_v1.font.size = Pt(11)
    p_rec.paragraph_format.space_after = Pt(4)
    
    p_date = doc.add_paragraph()
    r_l2 = p_date.add_run("DATE DE SORTIE : ")
    r_l2.bold = True
    r_l2.font.size = Pt(11)
    r_v2 = p_date.add_run(date_str)
    r_v2.font.size = Pt(11)
    p_date.paragraph_format.space_after = Pt(20)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_w = [Inches(4.2), Inches(1.1), Inches(1.7)]
    headers = ["Matériel", "Quantité", "Observations"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = col_w[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].bold = True
        set_cell_background(hdr_cells[i], "EFEFEF")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
    for it in items:
        row = table.add_row().cells
        row[0].text = it.get('name', '')
        row[0].width = col_w[0]
        set_cell_margins(row[0], top=100, bottom=100, left=150, right=150)
        
        row[1].text = str(it.get('quantity', 1))
        row[1].width = col_w[1]
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row[1], top=100, bottom=100, left=150, right=150)
        
        row[2].text = it.get('observations', '')
        row[2].width = col_w[2]
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row[2], top=100, bottom=100, left=150, right=150)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(45)
    
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sig = p_sig.add_run("La personne concernée")
    r_sig.bold = True
    r_sig.font.size = Pt(11)
    
    add_ismac_standard_footer(doc)
    return doc

# 3. DECHARGE DETAILLEE DU MATERIEL (Exact matching layout without table)
def generate_decharge_detaillee_exact(date_ville, designation_text, civilite, benef_name, pour_motif, en_date_du, num_serie_text):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_d = doc.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_d = p_d.add_run(date_ville or f"Rabat, le {datetime.datetime.now().strftime('%d/%m/%Y')}")
    r_d.font.size = Pt(11)
    p_d.paragraph_format.space_after = Pt(20)
    
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("BON DE DÉCHARGE\nDU MATÉRIEL")
    r_t.bold = True
    r_t.font.size = Pt(16)
    p_t.paragraph_format.space_after = Pt(26)
    
    lines = [
        ("DÉSIGNATION :  ", designation_text),
        (f"RÉCUPÉRÉ PAR  {civilite}. : ", benef_name.upper()),
        ("POUR : ", pour_motif),
        ("         EN DATE DU : ", en_date_du),
        ("         N° DE SERIE  :", f"{num_serie_text}          VU PAR ISMAC…………………………………………………………………")
    ]
    for lbl, val in lines:
        p = doc.add_paragraph()
        r_l = p.add_run(lbl)
        r_l.bold = True
        r_l.font.size = Pt(10.5)
        r_v = p.add_run(val)
        r_v.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(8)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(45)
    
    sig_t = doc.add_table(rows=1, cols=2)
    sig_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc = sig_t.rows[0].cells
    sc[0].width = Inches(3.5)
    sc[1].width = Inches(3.5)
    
    p_is = sc[0].paragraphs[0]
    p_is.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_is = p_is.add_run("Vu par ISMAC")
    r_is.bold = True
    r_is.font.size = Pt(11)
    
    p_cl = sc[1].paragraphs[0]
    p_cl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_cl = p_cl.add_run("La personne concernée")
    r_cl.bold = True
    r_cl.font.size = Pt(11)
    
    return doc

# 4. PROCES-VERBAL D'AFFECTATION MATERIEL
def generate_pv_affectation_exact(num_pv, objet_mouv, lieu_install, date_exec, resp_mise_en_place, items):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_t = doc.add_paragraph()
    r_t1 = p_t.add_run("PROCÈS-VERBAL D'AFFECTATION MATÉRIEL")
    r_t1.bold = True
    r_t1.font.size = Pt(13)
    r_t2 = p_t.add_run(f"N° : {num_pv or '01 / 2026'}")
    r_t2.bold = True
    r_t2.font.size = Pt(13)
    p_t.paragraph_format.space_after = Pt(12)
    
    fields = [
        ("Objet du mouvement : ", objet_mouv or "Installation et affectation des équipements audiovisuel"),
        ("Lieu d'installation : ", lieu_install or "Studios / Salles de cours ISMAC"),
        ("Date d'exécution : ", date_exec or "..........................................."),
        ("Responsable de la mise en place : ", resp_mise_en_place or "...........................................")
    ]
    for lbl, val in fields:
        p = doc.add_paragraph()
        r_l = p.add_run(lbl)
        r_l.bold = True
        r_l.font.size = Pt(10)
        r_v = p.add_run(val)
        r_v.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(2.8), Inches(1.6), Inches(0.9), Inches(1.7)]
    headers = ["Désignation du Matériel", "N° d'Inventaire", "Quantité", "Localisation"]
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = col_w[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], "EFEFEF")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        
    for it in items:
        row_c = table.add_row().cells
        row_c[0].text = it.get('name', '')
        row_c[0].width = col_w[0]
        set_cell_margins(row_c[0], top=80, bottom=80, left=120, right=120)
        
        row_c[1].text = it.get('inventory_number', 'ISMAC0001')
        row_c[1].width = col_w[1]
        row_c[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row_c[1], top=80, bottom=80, left=120, right=120)
        
        row_c[2].text = f"{it.get('quantity', 1):02d}" if isinstance(it.get('quantity', 1), int) else str(it.get('quantity', '01'))
        row_c[2].width = col_w[2]
        row_c[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row_c[2], top=80, bottom=80, left=120, right=120)
        
        row_c[3].text = it.get('location', lieu_install or '')
        row_c[3].width = col_w[3]
        set_cell_margins(row_c[3], top=80, bottom=80, left=120, right=120)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(35)
    
    sig_t = doc.add_table(rows=1, cols=3)
    sig_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_cells = sig_t.rows[0].cells
    labels_3 = ["Le magasinier\n........................................", "Responsable de la mise en place\n........................................", "Administration\n\n"]
    for i, lbl in enumerate(labels_3):
        p = s_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.font.size = Pt(10)
        r.bold = True
        
    add_ismac_standard_footer(doc)
    return doc

# 5. BON DE SORTIE DE MATERIEL
def generate_bon_sortie_exact(num_sortie, benef_name, motif_sortie, date_sortie, items):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_t = doc.add_paragraph()
    r_t1 = p_t.add_run("BON DE SORTIE DE MATÉRIEL")
    r_t1.bold = True
    r_t1.font.size = Pt(13)
    r_t2 = p_t.add_run(f"N° : {num_sortie or '01 / 2026'}")
    r_t2.bold = True
    r_t2.font.size = Pt(13)
    p_t.paragraph_format.space_after = Pt(12)
    
    fields = [
        ("Bénéficiaire : ", benef_name.upper()),
        ("Motif de la sortie : ", motif_sortie or "Affectation permanente"),
        ("Date de sortie du magasin : ", date_sortie or ".... / .... / 2026")
    ]
    for lbl, val in fields:
        p = doc.add_paragraph()
        r_l = p.add_run(lbl)
        r_l.bold = True
        r_l.font.size = Pt(10)
        r_v = p.add_run(val)
        r_v.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(2.8), Inches(1.6), Inches(0.9), Inches(1.7)]
    headers = ["Désignation précise du matériel", "N° d'Inventaire", "Quantité", "Observations / État"]
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].width = col_w[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9.5)
        set_cell_background(hdr_cells[i], "EFEFEF")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        
    for it in items:
        row_c = table.add_row().cells
        row_c[0].text = it.get('name', '')
        row_c[0].width = col_w[0]
        set_cell_margins(row_c[0], top=80, bottom=80, left=120, right=120)
        
        row_c[1].text = it.get('inventory_number', '')
        row_c[1].width = col_w[1]
        row_c[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row_c[1], top=80, bottom=80, left=120, right=120)
        
        row_c[2].text = f"{it.get('quantity', 1):02d}" if isinstance(it.get('quantity', 1), int) else str(it.get('quantity', '01'))
        row_c[2].width = col_w[2]
        row_c[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row_c[2], top=80, bottom=80, left=120, right=120)
        
        row_c[3].text = it.get('observations', '')
        row_c[3].width = col_w[3]
        set_cell_margins(row_c[3], top=80, bottom=80, left=120, right=120)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(35)
    
    sig_t = doc.add_table(rows=1, cols=2)
    sig_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_cells = sig_t.rows[0].cells
    labels_2 = ["Le Magasinier\n........................................", "Le Bénéficiaire\n........................................"]
    for i, lbl in enumerate(labels_2):
        p = s_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.font.size = Pt(10)
        r.bold = True
        
    add_ismac_standard_footer(doc)
    return doc
